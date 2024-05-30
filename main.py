'''
Be able to import word documents.
It will need to be able to adjust to a specifed font
It will need to be able to adjust to a specifed logo
It will need to be able to adjust to a specifed layout
Be able to export word documents accoridng to a template(This can be ignored)
'''
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

'''
Choose a logo
'''
def create_word_with_image(output_file, text, font_name='Arial', font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, image_path=None):
  # Create a new Document
  doc = Document()

  # Add a paragraph with the text
  paragraph = doc.add_paragraph(text)

  # Set the font name and size
  run = paragraph.runs[0]
  run.font.name = font_name
  run.font.size = Pt(font_size)

# Workaround to set the font name correctly
  r = run._element
  rFonts = r.find(qn('w:rFonts'))
  if rFonts is None:
     rFonts = OxmlElement('w:rFonts')
     r.insert(0, rFonts)
  rFonts.set(qn('w:eastAsia'), font_name)

  # Set the alignment
  paragraph.alignment = alignment

  # Add an image if provided
  if image_path:
      doc.add_picture(image_path)

  # Save the Word document
  doc.save(output_file)

# Example usage
create_word_with_image(
  'output.docx',
  'Hello, world!',
  font_name='Times New Roman',
  font_size=14,
  alignment=WD_ALIGN_PARAGRAPH.CENTER,
  image_path='example.jpg'  # Provide the path to your image file

'''
Convert from text to Word
'''

def convert_text_to_word(text_file, word_file):
    # Create a new Document
    doc = Document()

    # Open and read the text file
    with open(text_file, 'r') as file:
        content = file.read()

    # Add the text content to the Word document
    doc.add_paragraph(content)

    # Save the Word document
    doc.save(word_file)
    print(f"Converted {text_file} to {word_file}")

# Example usage
convert_text_to_word('example.txt', 'example.docx')